import yfinance as yf
import streamlit as st
import pandas as pd
import requests
import time
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Sidebar starts collapsed; Stock Screener page shows a clear instruction to open it ---
if "page" not in st.session_state:
    st.session_state.page = "🏛️ Congressional Trading"

st.set_page_config(page_title="Anita's Stock Screener", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""
    <style>
        .block-container {
            padding-top: 2rem;
            padding-bottom: 1rem;
        }
        html, body, [class*="css"] {
            font-size: 17px;
        }
        p, .stMarkdown, .stCaption, label {
            font-size: 17px !important;
        }
    </style>
""", unsafe_allow_html=True)

st.markdown("## Anita's Stock Screener")

# --- Version history (kept as a code comment for now — removed from UI since
#     the sidebar's mere presence prevents full collapse. Revisit placement later.) ---
# v6.6 — Jun 2026 — Condensed Filter Criteria into 3 columns instead of 2, removed redundant "Screen Stocks" heading/caption above the button, tightened top padding — goal: fit everything in one screen without scrolling
# v6.6 — Jun 2026 — Condensed Filter Criteria into a 4-column top row (Market/PE/ROE/Volume) plus a second row (Industry/Slider) — reduces from 3 rows to 2 rows so the Screen Stocks button is reachable with less scrolling
# v6.5 — Jun 2026 — Removed version history from sidebar entirely (its presence kept the sidebar from fully disappearing) — history now lives only as this code comment
# v6.4 — Jun 2026 — Moved Filter Criteria from sidebar to a visible expander directly on the Stock Screener page (avoids the hidden-sidebar-arrow problem entirely) — filters auto-show when Stock Screener is selected, disappear when switching to Congressional Trading, and values are retained if the user switches back
# v6.3 — Jun 2026 — Removed unreliable JS auto-open attempt (Streamlit limitation: sidebar state can't reliably auto-toggle per page) — replaced with a clear, always-visible instruction on the Stock Screener page instead
# v6.2 — Jun 2026 — Added JS attempt to auto-open sidebar on Stock Screener page, plus a reliable visible instruction as backup since JS may not always work
# v6.1 — Jun 2026 — Sidebar now opens automatically on Stock Screener page and stays hidden on Congressional Trading page | Removed redundant grey button — single working blue button
# v6.0 — Jun 2026 — Compact layout: reduced padding, smaller title, removed divider, condensed Congressional Trading instructions, version history moved to small sidebar icon
# v5.4 — Jun 2026 — Congressional Trading moved to top of menu | Removed ticker box (was misleading — real filtering happens on Capitol Trades) | Added roadmap of what to do once on Capitol Trades | Capitol Trades now opens in a new tab so the app is never replaced — just close that tab to return
# v5.3 — Jun 2026 — Added friendly message for Yahoo Finance rate limiting
# v5.2 — Jun 2026 — Simplified Congressional Trading — direct Capitol Trades redirect | v5.1 — Jun 2026 — Fixed Capitol Trades URL | v5.0 — Jun 2026 — Removed Google Sheets button | Hidden version history
# v4.x — Switched Congressional Trading to Senate eFD live search | Fixed Senate data parsing | Fixed Senate Stock Watcher URL | Added Congressional Trading Tracker
# v3.0 — Jun 2026 — Switched to Yahoo Finance
# v2.0 — Jun 2026 — Fixed FMP endpoints
# v1.9 — May 2026 — Initial release

# --- Navigation ---
page = st.radio(
    "Select a tool:",
    ["🏛️ Congressional Trading", "📈 Stock Screener"],
    horizontal=True,
    index=["🏛️ Congressional Trading", "📈 Stock Screener"].index(st.session_state.page)
)

if page != st.session_state.page:
    st.session_state.page = page
    st.rerun()

# ===========================================================================
# PAGE 1 — STOCK SCREENER
# ===========================================================================
if page == "📈 Stock Screener":

    st.warning("""
    ⚠️ **This feature is currently unavailable.**
    
    We are working on improving the Stock Screener and it will be fully demonstrated in **Video 2** — coming soon!
    
    In the meantime, please enjoy the **🏛️ Congressional Trading Tracker** on the left menu!
    """)
    st.stop()


    TICKERS = {
        "S&P 500": [
            "AAPL","MSFT","NVDA","AMZN","META","GOOGL","GOOG","BRK-B","LLY","AVGO",
            "JPM","TSLA","UNH","V","XOM","MA","JNJ","PG","COST","HD","MRK","ABBV",
            "CVX","BAC","CRM","NFLX","AMD","PEP","KO","TMO","ACN","MCD","CSCO","WMT",
            "ABT","DHR","ADBE","TXN","LIN","PM","NEE","ORCL","RTX","QCOM","MS","HON",
            "UPS","AMGN","IBM","CAT","GS","BA","SBUX","GE","INTU","SPGI","BLK","AXP",
            "ELV","MDT","DE","GILD","ADI","VRTX","ISRG","MMC","SYK","REGN","ZTS","PLD",
            "CI","CB","ADP","SCHW","C","MO","SO","DUK","TGT","CME","EOG","SLB","BDX",
            "ITW","BSX","NOC","WM","AON","PNC","USB","MMM","FDX","EMR","APD","CL","NSC"
        ],
        "Nasdaq 100": [
            "AAPL","MSFT","NVDA","AMZN","META","GOOGL","GOOG","TSLA","AVGO","ASML",
            "COST","NFLX","AMD","QCOM","ADBE","INTU","TXN","AMGN","PEP","CSCO",
            "ISRG","AMAT","HON","VRTX","REGN","MU","LRCX","ADI","KLAC","PANW",
            "MELI","SNPS","CDNS","CRWD","CEG","FTNT","ORLY","MNST","CTAS","PAYX",
            "MRVL","KDP","AEP","DXCM","ODFL","ROST","FAST","CTSH","TEAM","IDXX",
            "BIIB","ILMN","NXPI","WDAY","PCAR","VRSK","DDOG","ZS","ANSS","ALGN",
            "WBD","FANG","GEHC","GFS","ON","EXC","XEL","SIRI","DLTR","ENPH","LCID"
        ],
        "Dow Jones (DJIA)": [
            "AAPL","AMGN","AXP","BA","CAT","CRM","CSCO","CVX","DIS","DOW",
            "GS","HD","HON","IBM","JNJ","JPM","KO","MCD","MMM","MRK",
            "MSFT","NKE","PG","TRV","UNH","V","VZ","WBA","WMT","RTX"
        ],
        "Russell 2000 (Sample)": [
            "SAIA","WING","BOOT","FORM","MGNI","PRGS","ITRI","ENVA","CSWI","BCPC",
            "MGEE","LANC","NWBI","SFBS","TOWN","WAFD","CVBF","NBTB","FULT","WSFS",
            "CAKE","TXRH","BJRI","RRGB","DENN","CHUY","FRSH","HURN","AMSF","PRIM",
            "ARCB","HTLD","MRTN","MATX","HUBG","JBHT","WERN","KNX",
            "AGIO","ACAD","FOLD","HRMY","PTGX","RCUS","SERA","XNCR","IGMS","KROS",
            "ADUS","AMED","FWRG","LHCG","MGLN","PDCO","PRSC","QTWO","RCKY","RELY"
        ]
    }

    # --- Filter Criteria — shown on main page inside an expander ---
    with st.expander("⚙️ **Set Filter Criteria**", expanded=True):
        col_a, col_b, col_c, col_d = st.columns(4)
        with col_a:
            market = st.selectbox("Market Index", list(TICKERS.keys()))
        with col_b:
            pe_max = st.number_input("Max PE Ratio", min_value=0.0, value=25.0)
        with col_c:
            roe_min = st.number_input("Min ROE (%)", min_value=0.0, value=10.0)
        with col_d:
            volume_min = st.number_input("Min Volume", min_value=0, value=1000000)

        col_e, col_f = st.columns([1, 2])
        with col_e:
            industry = st.text_input("Industry (optional)", value="")
        with col_f:
            all_symbols = TICKERS[market]
            max_stocks = st.slider(
                "Max stocks to screen", min_value=10, max_value=len(all_symbols),
                value=min(30, len(all_symbols)), step=10
            )

    # --- Helper: Build HTML table ---
    def build_html(df):
        rows = ""
        for _, row in df.iterrows():
            cells = ""
            for col, v in zip(df.columns, row):
                if col == "Yahoo Finance":
                    cells += f'<td><a href="{v}" target="_blank">View</a></td>'
                else:
                    cells += f"<td>{v}</td>"
            rows += f"<tr>{cells}</tr>\n"
        headers = "".join(f"<th>{c}</th>" for c in df.columns)
        return f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <title>Stock Screener Results</title>
  <style>
    body {{ font-family: Arial, sans-serif; padding: 20px; }}
    h2 {{ color: #333; }}
    table {{ border-collapse: collapse; width: 100%; }}
    th {{ background-color: #4CAF50; color: white; padding: 8px 12px; text-align: left; }}
    td {{ border: 1px solid #ddd; padding: 8px 12px; }}
    tr:nth-child(even) {{ background-color: #f9f9f9; }}
    tr:hover {{ background-color: #f1f1f1; }}
  </style>
</head>
<body>
  <h2>Anita's Stock Screener Results</h2>
  <table>
    <thead><tr>{headers}</tr></thead>
    <tbody>{rows}</tbody>
  </table>
</body>
</html>"""

    # --- Helper: Build Word (.docx) ---
    def build_docx(df):
        doc = Document()
        title = doc.add_heading("Anita's Stock Screener Results", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        doc.add_paragraph("")
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = col
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "4CAF50")
            tcPr.append(shd)
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        for col in table.columns:
            for cell in col.cells:
                cell.width = Inches(1.2)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    # --- Core data fetch using Yahoo Finance ---
    def get_stock_data(symbol):
        try:
            ticker = yf.Ticker(symbol)
            info   = ticker.info
            if not info:
                return None, "No data returned from Yahoo Finance"
            price = info.get("currentPrice") or info.get("regularMarketPrice")
            pe    = info.get("trailingPE")
            vol   = info.get("averageVolume")
            name  = info.get("longName", symbol)
            ind   = info.get("industry", "")
            roe   = info.get("returnOnEquity")
            if roe is not None:
                roe = round(float(roe) * 100, 2)
            if price is None:
                return None, "No price data — symbol may be delisted or invalid"
            return {"pe": pe, "roe": roe, "vol": vol, "ind": ind, "name": name, "price": price}, None
        except Exception as e:
            err_text = str(e)
            if "Too Many Requests" in err_text or "Rate limited" in err_text or "429" in err_text:
                return None, "RATE_LIMITED"
            return None, f"Error: {err_text}"

    # --- Main Screener ---
    if st.button("🚀 Screen Stocks"):
        symbols  = all_symbols[:max_stocks]
        st.info(f"Screening {len(symbols)} stocks from {market}… this may take a minute, please wait!")
        results  = []
        skipped  = []
        skip_log = []
        rate_limited_count = 0
        progress = st.progress(0)
        status   = st.empty()

        for i, symbol in enumerate(symbols):
            status.text(f"Checking {symbol} ({i+1}/{len(symbols)})…")
            data, error = get_stock_data(symbol)
            if not data:
                skipped.append(symbol)
                if error == "RATE_LIMITED":
                    rate_limited_count += 1
                    skip_log.append(f"**{symbol}** — Temporarily unavailable (high traffic)")
                else:
                    skip_log.append(f"**{symbol}** — {error}")
                progress.progress((i + 1) / len(symbols))
                time.sleep(0.2)
                continue
            pe, roe, vol, ind, name, price = data["pe"], data["roe"], data["vol"], data["ind"], data["name"], data["price"]
            if pe is None and roe is None:
                skipped.append(symbol)
                skip_log.append(f"**{symbol}** — No PE or ROE data available")
                progress.progress((i + 1) / len(symbols))
                time.sleep(0.2)
                continue
            pe_ok  = pe  is not None and pe  <= pe_max
            roe_ok = roe is not None and roe >= roe_min
            vol_ok = vol is not None and vol >= volume_min
            ind_ok = industry.lower() in ind.lower() if industry else True
            if pe_ok and roe_ok and vol_ok and ind_ok:
                results.append({
                    "Symbol":     symbol,
                    "Company":    name,
                    "Price":      f"${price:.2f}" if price else "N/A",
                    "PE Ratio":   f"{pe:.2f}"     if pe    else "N/A",
                    "ROE (%)":    f"{roe:.2f}"    if roe   else "N/A",
                    "Avg Volume": f"{vol:,}"       if vol   else "N/A",
                    "Industry":   ind,
                })
            else:
                reasons = []
                if not pe_ok:  reasons.append(f"PE={round(pe,1) if pe else 'N/A'} (max {pe_max})")
                if not roe_ok: reasons.append(f"ROE={round(roe,1) if roe else 'N/A'}% (min {roe_min}%)")
                if not vol_ok: reasons.append("Vol too low")
                if not ind_ok: reasons.append(f"Industry='{ind}'")
                skip_log.append(f"**{symbol}** — filtered out: {', '.join(reasons)}")
            progress.progress((i + 1) / len(symbols))
            time.sleep(0.3)

        status.empty()
        progress.empty()

        # --- Friendly message if rate limiting affected most of the screen ---
        if rate_limited_count > 0 and rate_limited_count >= len(symbols) * 0.5:
            st.warning(
                "🚦 **Yahoo Finance is experiencing high traffic right now.**\n\n"
                "This isn't a problem with the app — Yahoo's free data service is temporarily "
                "limiting requests from many users at once. This usually clears up within a "
                "few minutes.\n\n"
                "**What you can try:**\n"
                "- Wait 5–10 minutes and click 🚀 Screen Stocks again\n"
                "- Try a smaller index like Dow Jones, or reduce the number of stocks to screen\n\n"
                "💡 *I'm aware of this limitation and am exploring a more reliable, paid data "
                "source for the future so this happens less often. Thank you for your patience "
                "as we improve the app together!* — Anita"
            )

        if results:
            st.success(f"✅ Found {len(results)} matching stocks from {market}!")
            df = pd.DataFrame(results)
            df["Yahoo Finance"] = df["Symbol"].apply(lambda s: f"https://finance.yahoo.com/quote/{s}")
            st.dataframe(df, column_config={"Yahoo Finance": st.column_config.LinkColumn("🔗 More Info", help="Click to open on Yahoo Finance", display_text="View")}, hide_index=True)

            st.markdown("### 📥 Download Results")
            st.caption("Choose the format that works best for you:")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.download_button("📄 CSV", df.to_csv(index=False), "screener_results.csv", "text/csv")
                st.caption("Works in Excel, Google Sheets, Numbers")
            with col2:
                st.download_button("🌐 HTML Table", build_html(df), "screener_results.html", "text/html")
                st.caption("Clickable links in any web browser")
            with col3:
                st.download_button("📝 Word Doc", build_docx(df), "screener_results.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.caption("Opens in Word, Google Docs, LibreOffice")
        else:
            if rate_limited_count < len(symbols) * 0.5:
                st.warning("No stocks matched your criteria. Try relaxing your filters!")

        if skip_log:
            with st.expander(f"ℹ️ Details — {len(skipped)} skipped + filtered stocks", expanded=False):
                st.caption("Skipped = no data. Filtered out = data found but didn't meet your criteria.")
                for line in skip_log:
                    st.markdown(line)


# ===========================================================================
# PAGE 2 — CONGRESSIONAL TRADING TRACKER
# ===========================================================================
elif page == "🏛️ Congressional Trading":

    st.markdown("### 🏛️ Congressional Trading Tracker")
    st.caption("Explore US Congress stock trades (Senate + House) — powered by Capitol Trades")

    st.markdown(
        "**📋 Click below to open Capitol Trades**, then use its filters to narrow by ticker, "
        "politician, party, or trade size. Quick ideas: 🔍 search a company · 🟢 filter **Buy** only · "
        "💰 sort by trade size · 🧑‍⚖️ follow a politician — trades are sorted by most recent by default."
    )

    capitol_trades_url = "https://www.capitoltrades.com/trades"
    st.markdown(
        f'<a href="{capitol_trades_url}" target="_blank">'
        f'<button style="background-color:#1a73e8;color:white;border:none;padding:12px 24px;'
        f'border-radius:6px;cursor:pointer;font-size:17px;width:100%;margin-top:10px">'
        f'🏛️ Open Capitol Trades (opens in a new tab)</button></a>',
        unsafe_allow_html=True
    )
    st.caption(
        "💡 Capitol Trades opens in a **new browser tab** — your Stock Screener app stays "
        "open here. When you're done exploring, just **close that tab** to come back!"
    )

    st.markdown("---")
    st.markdown(
        "📌 **About Capitol Trades:** A free, independent website tracking all US Congress "
        "stock trades filed under the STOCK Act. Covers both Senate and House of Representatives. "
        "Updated daily with the latest disclosures.\n\n"
        "[Visit Capitol Trades](https://www.capitoltrades.com) | "
        "[Official Senate Filings](https://efdsearch.senate.gov) | "
        "[Official House Filings](https://disclosures-clerk.house.gov/FinancialDisclosure)"
    )
